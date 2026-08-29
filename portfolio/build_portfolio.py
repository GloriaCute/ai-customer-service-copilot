from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"E:\Gloria\ready\portfolio")
ROOT.mkdir(parents=True, exist_ok=True)
FONT = r"C:\Windows\Fonts\NotoSansSC-VF.ttf"
BLUE = "#1F5EFF"; NAVY = "#102A56"; SKY = "#EAF2FF"; BG = "#F7F9FC"
INK = "#172B4D"; MUTED = "#637083"; LINE = "#DDE4EE"; GREEN = "#0F9D6A"; ORANGE = "#E57A15"; RED = "#C2414B"

def f(size, bold=False):
    return ImageFont.truetype(FONT, size, layout_engine=ImageFont.Layout.BASIC)

def wrap(draw, text, font, width):
    lines=[]
    for part in text.split("\n"):
        line=""
        for ch in part:
            candidate=line+ch
            if draw.textlength(candidate, font=font) <= width:
                line=candidate
            else:
                if line: lines.append(line)
                line=ch
        lines.append(line)
    return lines

def text(draw, xy, value, size=28, fill=INK, maxw=None, gap=10, bold=False):
    font=f(size, bold)
    lines=wrap(draw, value, font, maxw) if maxw else value.split("\n")
    x,y=xy
    for line in lines:
        draw.text((x,y), line, font=font, fill=fill)
        y += size+gap
    return y

def rounded(draw, box, fill="white", outline=None, radius=20, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)

def header(draw, kicker, title, sub, W=1600):
    draw.rectangle((0,0,W,118), fill="white")
    draw.line((60,117,W-60,117), fill=LINE, width=2)
    text(draw,(60,24),kicker.upper(),18,BLUE,bold=True)
    text(draw,(60,51),title,35,NAVY,bold=True)
    text(draw,(W-370,60),sub,17,MUTED,maxw=300)

def node(draw, x,y,w,h,label,sub="",color=BLUE):
    rounded(draw,(x,y,x+w,y+h),"white",color,22,3)
    draw.ellipse((x+18,y+22,x+58,y+62),fill=color)
    text(draw,(x+75,y+20),label,23,INK,maxw=w-95,bold=True)
    if sub: text(draw,(x+75,y+55),sub,16,MUTED,maxw=w-95)

def arrow(draw, p1,p2,color=BLUE):
    draw.line((p1,p2),fill=color,width=5)
    x,y=p2; draw.polygon([(x,y),(x-16,y-9),(x-16,y+9)],fill=color)

def save(img, name):
    path=ROOT/name; img.save(path, quality=95); return path

def architecture_initial():
    W,H=1600,900; img=Image.new("RGB",(W,H),BG); d=ImageDraw.Draw(img); header(d,"Architecture 01","最初的 RAG 架构","实验 1 前的基础链路",W)
    items=[("用户问题","原始 Query"),("知识检索","Chunk · Embedding · Hybrid Search\nRerank · Top K"),("最终 LLM","读取 Query 与 Context"),("客服建议回复","回答 / 引用 / 拒答")]
    xs=[80,445,875,1245]
    for i,(a,b) in enumerate(items):
        node(d,xs[i],350,260,150,a,b)
        if i<3: arrow(d,(xs[i]+260,425),(xs[i+1]-20,425))
    rounded(d,(80,615,1520,800),"#EFF5FF",None,24)
    text(d,(115,650),"这一阶段的关键风险",25,NAVY,bold=True)
    text(d,(115,700),"原始 PDF 的 Chunk 可能过碎或过大；检索系统会在已有资料中选“最相似”的内容，因此知识库外问题也可能得到无关召回。",22,INK,maxw=1320)
    return save(img,"01_最初RAG架构.png")

def architecture_v21():
    W,H=1600,980; img=Image.new("RGB",(W,H),BG); d=ImageDraw.Draw(img); header(d,"Architecture 02","实验 3：复合 Query 的多路检索","V2.1 / Query Decomposition",W)
    node(d,70,350,220,120,"用户问题","原始 query")
    node(d,380,330,270,160,"查询拆解 LLM","输出 query_1 / query_2")
    node(d,770,245,270,125,"知识检索 A","query_1 → result_1",GREEN)
    node(d,770,475,270,125,"知识检索 B","query_2 → result_2",GREEN)
    node(d,1150,350,260,135,"合并检索结果","merged_result",ORANGE)
    node(d,1160,615,260,135,"最终 LLM","query + merged_result",BLUE)
    node(d,1160,835,260,100,"回答","引用 / 拒答 / 建议回复",NAVY)
    arrow(d,(290,410),(360,410)); arrow(d,(650,390),(745,305)); arrow(d,(650,430),(745,535)); arrow(d,(1040,305),(1130,390)); arrow(d,(1040,535),(1130,445)); arrow(d,(1280,485),(1280,595)); arrow(d,(1290,750),(1290,815))
    rounded(d,(80,125,1510,205),"#EAF2FF",None,18)
    text(d,(110,145),"目标：把“会员 + 家电”这类复合问题拆成独立子问题，分别检索后再合并证据。",23,NAVY,maxw=1320,bold=True)
    return save(img,"02_V21多路检索架构.png")

def experiments():
    W,H=1600,920; img=Image.new("RGB",(W,H),BG); d=ImageDraw.Draw(img); header(d,"Experiment Review","从文档问题到多路检索","三个 PoC 实验",W)
    cols=[(70,"实验 1｜Chunk 粒度问题","原始 PDF 直接入库","Chunk 过碎或 700–1000 字；多个主题、网页导航与 URL 混在一起。","发现：语义完整性与噪声需要平衡。",RED),(545,"实验 2｜FAQ 级 Chunk","清洗 Markdown · FAQ 级 Chunk","单一意图更容易命中对应规则；会员卡、手机号、重量限制等检索更清晰。","代价：复合问题需要同时召回多张知识卡。",ORANGE),(1020,"实验 3｜多路检索","查询拆解 · 双路检索","将复合问题拆为 query_1 / query_2，合并两路证据后交给最终 LLM。","结果：复合问题改善；延迟和调用成本增加。",GREEN)]
    for x,title,tag,body,out,color in cols:
        rounded(d,(x,210,x+410,790),"white",LINE,24)
        d.rectangle((x,210,x+410,222),fill=color)
        text(d,(x+28,255),title,26,NAVY,maxw=350,bold=True)
        rounded(d,(x+28,335,x+365,390),"#F2F6FC",None,14); text(d,(x+45,350),tag,18,color,maxw=300,bold=True)
        text(d,(x+28,435),body,22,INK,maxw=350)
        d.line((x+28,620,x+365,620),fill=LINE,width=2)
        text(d,(x+28,650),out,20,NAVY,maxw=350,bold=True)
    text(d,(90,835),"结论：不是“正式版本升级”，而是在 PoC 中围绕问题、假设和证据开展的三轮实验。",22,MUTED,maxw=1380)
    return save(img,"03_三轮实验复盘.png")

def diagnosis():
    W,H=1600,970; img=Image.new("RGB",(W,H),BG); d=ImageDraw.Draw(img); header(d,"Diagnosis Method","回答错误时，先定位在哪一层出问题","不要直接归因于模型",W)
    steps=[("回答错误","先记录原始问题、回答与证据"),("检查 Retrieval","正确 Chunk 是否进入 Top K？"),("未召回","检查 Query / Chunk / Embedding / Top K / Ranking"),("已召回","检查 Context / Prompt / LLM / Answer 输出")]
    y=190
    for i,(a,b) in enumerate(steps):
        node(d,150,y,480,105,a,b,BLUE if i<2 else (ORANGE if i==2 else GREEN))
        if i<3: arrow(d,(390,y+105),(390,y+135))
        y+=150
    rounded(d,(800,190,1510,820),"white",LINE,24)
    text(d,(835,225),"实际遇到的 Bug",26,NAVY,bold=True)
    bugs=["Query 没有传给最终 LLM","Retrieval 结果未真正插入 Prompt","Answer 节点输出固定文案","Chunk 过大 / 过碎","单次检索漏掉复合规则","API quota exceeded"]
    yy=295
    for b in bugs:
        d.ellipse((840,yy+7,855,yy+22),fill=ORANGE); text(d,(875,yy),b,21,INK,maxw=560); yy+=72
    rounded(d,(150,820,630,900),"#EAF2FF",None,18)
    text(d,(180,842),"Top K 3 → 5 仍未召回第二条规则",20,NAVY,maxw=410,bold=True)
    return save(img,"04_问题定位方法.png")

def app_frame(title):
    W,H=1440,900; img=Image.new("RGB",(W,H),"#F5F7FB"); d=ImageDraw.Draw(img)
    d.rectangle((0,0,W,78),fill="white"); text(d,(45,24),"企业 AI 客服知识助手",27,NAVY,bold=True); text(d,(400,32),"RAG Customer Service Copilot Demo",16,MUTED); rounded(d,(1170,20,1360,55),"#E8F8F0",None,18); text(d,(1200,28),"知识库已连接",15,GREEN,bold=True)
    text(d,(45,110),title,26,NAVY,bold=True); return img,d

def prototype_workbench():
    img,d=app_frame("客服工作台｜客户 A · 待处理")
    rounded(d,(35,155,345,850),"white",LINE,18); text(d,(60,185),"客户会话",20,NAVY,bold=True)
    people=[("客户 A","我是会员，半年前买了一台家电，可以退吗？","待处理"),("客户 B","会员卡丢了怎么办？","AI 已生成"),("客户 C","35 公斤商品能包裹配送吗？","已完成")]
    y=235
    for name,q,st in people:
        rounded(d,(55,y,325,y+130),"#F6F9FD" if name=="客户 A" else "white",LINE,14); text(d,(75,y+16),name,19,NAVY,bold=True); text(d,(75,y+48),q,15,INK,maxw=220); text(d,(75,y+100),st,14,ORANGE if st=="待处理" else GREEN); y+=150
    rounded(d,(370,155,910,850),"white",LINE,18); text(d,(400,190),"客户对话",20,NAVY,bold=True); rounded(d,(420,270,835,350),"#EAF2FF",None,18); text(d,(445,290),people[0][1],19,INK,maxw=360)
    rounded(d,(400,700,870,800),"#F6F8FC",LINE,15); text(d,(430,730),"输入或编辑客服回复…",18,MUTED); rounded(d,(705,815,870,855),BLUE,None,16); text(d,(733,824),"发送",16,"white",bold=True)
    rounded(d,(935,155,1405,850),"white",LINE,18); text(d,(965,190),"AI Copilot",22,NAVY,bold=True); text(d,(965,240),"建议回复",16,MUTED)
    answer="家用电器不适用于退货政策，因此即使是会员且购买时间在 365 天内，也不能仅按该政策办理退货。"
    text(d,(965,275),answer,21,INK,maxw=390)
    rounded(d,(965,450,1375,560),"#F2F6FC",None,15); text(d,(990,472),"知识依据",17,NAVY,bold=True); text(d,(990,505),"• 会员 365 天退货规则\n• 家用电器属于排除范围",16,INK,maxw=350)
    for i,label in enumerate(["采用回复","编辑","重新生成","转人工"]):
        x=965+(i%2)*205; y=610+(i//2)*62; rounded(d,(x,y,x+185,y+45),"#EAF2FF" if i<2 else "white",BLUE if i<2 else LINE,12); text(d,(x+25,y+11),label,16,BLUE if i<2 else NAVY,bold=True)
    return save(img,"05_产品原型_客服工作台.png")

def prototype_sources():
    img,d=app_frame("AI 建议回复｜有依据的客服答复")
    rounded(d,(50,160,870,830),"white",LINE,20); text(d,(85,205),"客户问题",18,MUTED); text(d,(85,242),"会员卡丢了怎么办？",28,NAVY,bold=True); d.line((85,300,830,300),fill=LINE,width=2); text(d,(85,330),"AI 建议回复",18,MUTED)
    text(d,(85,370),"可以在宜家商场的会员自助机器上申请遗失会员卡补领，补领不会产生任何费用。",26,INK,maxw=690)
    rounded(d,(85,535,830,640),"#EAF8F1",None,18); text(d,(115,558),"建议状态：可采用",19,GREEN,bold=True); text(d,(115,592),"已匹配 1 条直接知识依据",17,INK)
    for i,label in enumerate(["采用回复","编辑后采用","重新生成"]):
        x=85+i*205; rounded(d,(x,690,x+180,745),BLUE if i==0 else "white",BLUE,15); text(d,(x+28,705),label,17,"white" if i==0 else BLUE,bold=True)
    rounded(d,(910,160,1390,830),"white",LINE,20); text(d,(945,205),"知识依据",22,NAVY,bold=True); rounded(d,(945,265,1355,400),"#F5F8FC",LINE,14); text(d,(970,288),"宜家会员 FAQ",18,NAVY,bold=True); text(d,(970,323),"会员卡丢失了怎么办？\n会员自助机器可申请补领，且不收费。",17,INK,maxw=350)
    rounded(d,(945,440,1355,560),"#F6F8FC",None,14); text(d,(970,465),"AI 处理详情",18,NAVY,bold=True); text(d,(970,500),"状态：成功\n查询拆解：否\n引用文档：1",16,INK)
    return save(img,"06_产品原型_知识依据.png")

def prototype_escalation():
    img,d=app_frame("知识不足｜拒答与人工确认")
    rounded(d,(90,170,1350,785),"white",LINE,22); text(d,(135,220),"客户问题",18,MUTED); text(d,(135,255),"南京门店今天几点关门？",29,NAVY,bold=True)
    rounded(d,(135,355,1305,545),"#FFF8EC",None,18); text(d,(175,392),"知识依据不足，建议人工确认",25,ORANGE,bold=True); text(d,(175,440),"根据当前知识库无法确定。当前资料只包含退货、送货和会员 FAQ，未覆盖门店营业时间。",22,INK,maxw=1030)
    rounded(d,(135,610,390,670),BLUE,None,16); text(d,(190,627),"转人工确认",19,"white",bold=True); rounded(d,(420,610,660,670),"white",BLUE,16); text(d,(475,627),"查看知识范围",19,BLUE,bold=True)
    rounded(d,(910,610,1305,710),"#F4F7FC",None,15); text(d,(940,630),"不会凭模型自身知识补充营业时间",17,NAVY,bold=True); text(d,(940,662),"人工是最终决策者",16,MUTED)
    return save(img,"07_产品原型_拒答转人工.png")

def prototype_admin():
    img,d=app_frame("主管视角｜测试与效果管理")
    for x,title,desc in [(50,"测试集","固定题目 · 同题对比"),(400,"异常案例","复合规则 / 知识库外"),(750,"待验证指标","采纳率 · 修改率 · 时长")]:
        rounded(d,(x,155,x+300,285),"white",LINE,16); text(d,(x+25,180),title,19,NAVY,bold=True); text(d,(x+25,220),desc,16,MUTED,maxw=240)
    rounded(d,(50,330,900,820),"white",LINE,18); text(d,(80,360),"实验观察",22,NAVY,bold=True)
    rows=[("单一意图","V2：FAQ 级 Chunk 更精准","已验证"),("复合规则","V2：证据可能不完整","发现问题"),("Query Decomposition","V2.1：两题复合问题成功","已验证"),("延迟","约 14 秒 / 17 秒","需优化")]
    y=415
    for a,b,c in rows:
        d.line((80,y+58,870,y+58),fill=LINE,width=1); text(d,(90,y),a,17,NAVY,bold=True); text(d,(285,y),b,17,INK,maxw=355); text(d,(700,y),c,16,GREEN if c=="已验证" else ORANGE); y+=72
    rounded(d,(935,330,1390,820),"white",LINE,18); text(d,(965,360),"上线后待验证",22,NAVY,bold=True)
    metrics=["AI 建议回复采纳率","人工修改率","平均处理时长","错误回答率","正确拒答率","单次调用成本"]
    y=420
    for m in metrics:
        rounded(d,(965,y,1355,y+45),"#F6F8FC",None,12); text(d,(985,y+11),m,16,INK); text(d,(1265,y+11),"待验证",15,MUTED); y+=55
    return save(img,"08_产品原型_主管后台.png")

def long_image(paths):
    W,H=1440,10400; img=Image.new("RGB",(W,H),BG); d=ImageDraw.Draw(img)
    y=0
    d.rectangle((0,0,W,720),fill=NAVY); text(d,(100,130),"企业 AI 客服知识助手",54,"white",bold=True); text(d,(100,215),"RAG PoC｜从问题定位到多路检索验证",31,"#CFE0FF"); text(d,(100,325),"个人学习 PoC｜示例知识来自公开客服资料｜非宜家官方项目",20,"#A9C4F5"); text(d,(100,510),"[姓名占位]  ·  AI 产品实习作品集",22,"white"); y=800
    sections=[
        ("01  用户、角色与业务问题","使用者：一线客服\n管理者：客服主管 / 运营负责人\n决策者：客服中心负责人 / 数字化负责人\n价值假设：减少查政策时间、降低规则判断错误、提高处理效率",None),
        ("02  产品方案与核心交互","客户提问 → AI 检索企业知识 → 生成建议回复与知识依据 → 客服采用、编辑、重新生成或转人工。",paths["workbench"]),
        ("03  最初的 RAG 架构","User Query → Knowledge Retrieval → LLM → Answer。基础链路跑通后，重点转向 Chunk、检索质量与拒答机制。",paths["initial"]),
        ("04  实验 1 与实验 2","原始 PDF 的 Chunk 过碎或过大；清洗为 FAQ 级 Chunk 后，单一意图检索更精准，但复合规则需要同时命中多张知识卡。",paths["experiments"]),
        ("05  测试方法与复合问题失败","固定测试集 → 同样问题 → 比较 Top1 / Top3 / 噪声 / 最终回答 → 记录异常案例。Top K 3→5 后仍遗漏第二条必要规则。",paths["diagnosis"]),
        ("06  实验 3：Query Decomposition","将复杂问题拆成 query_1 / query_2，分别检索，再合并 merged_result 交给最终 LLM。",paths["v21"]),
        ("07  结果与 Trade-off","两类复合问题均获得正确结论；观察到约 14 秒、17 秒延迟。效果提升伴随额外 LLM 调用和检索成本。",None),
        ("08  产品原型与价值","AI 辅助人工而非替代人工。上线后需验证采纳率、人工修改率、处理时长、错误回答率、正确拒答率与单次调用成本。",paths["admin"]),
        ("09  限制与下一步","当前为个人 PoC：无真实客服、生产流量与 ROI 数据。下一步验证真实使用流程、延迟优化和人工反馈闭环。",paths["escalation"]),
    ]
    for title,body,pic in sections:
        text(d,(90,y),title,33,NAVY,bold=True); y+=60; text(d,(90,y),body,22,INK,maxw=1230); y+=145
        if pic:
            im=Image.open(pic); im.thumbnail((1150,680)); x=(W-im.width)//2; img.paste(im,(x,y)); y+=im.height+80
        else:
            rounded(d,(90,y,1350,y+38),"#EAF2FF",None,16); y+=95
    # 根据实际内容裁切底部空白，保留适量收尾留白。
    img = img.crop((0, 0, W, min(H, y + 120)))
    return save(img,"企业AI客服RAG_PoC_作品集长图.png")

def set_cell_shading(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def set_run(run, size=11, color="172B4D", bold=False):
    run.font.name="Microsoft YaHei"; run._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); run.font.size=Pt(size); run.font.color.rgb=RGBColor.from_string(color); run.bold=bold

def add_p(doc, value="", size=11, color="172B4D", bold=False, style=None, align=None, after=6):
    p=doc.add_paragraph(style=style); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.10
    if align is not None: p.alignment=align
    r=p.add_run(value); set_run(r,size,color,bold); return p

def add_bullets(doc, values):
    for v in values:
        p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.10; set_run(p.add_run(v),11)

def add_pic(doc,path,width=6.2):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(path),width=Inches(width)); p.paragraph_format.space_after=Pt(8)

def page_break(doc): doc.add_page_break()

def docx(paths):
    doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(1); sec.bottom_margin=Inches(1); sec.left_margin=Inches(1); sec.right_margin=Inches(1); sec.header_distance=Inches(.492); sec.footer_distance=Inches(.492)
    styles=doc.styles
    styles['Normal'].font.name='Microsoft YaHei'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); styles['Normal'].font.size=Pt(11); styles['Normal'].paragraph_format.space_after=Pt(6); styles['Normal'].paragraph_format.line_spacing=1.10
    for name,size,color,before,after in [('Heading 1',16,'2E74B5',16,8),('Heading 2',13,'2E74B5',12,6),('Heading 3',12,'1F4D78',8,4)]:
        st=styles[name]; st.font.name='Microsoft YaHei'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.font.bold=True; st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after)
    h=sec.header.paragraphs[0]; h.alignment=WD_ALIGN_PARAGRAPH.RIGHT; set_run(h.add_run('企业 AI 客服 RAG PoC｜作品集'),9,'637083')
    ft=sec.footer.paragraphs[0]; ft.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_run(ft.add_run('个人学习 PoC｜示例知识来自公开客服资料｜非宜家官方项目'),8,'637083')
    # cover
    for _ in range(5): add_p(doc,'',after=10)
    add_p(doc,'AI 产品实习作品集',14,'1F5EFF',True,align=WD_ALIGN_PARAGRAPH.CENTER,after=10)
    add_p(doc,'企业 AI 客服知识助手',30,'102A56',True,align=WD_ALIGN_PARAGRAPH.CENTER,after=6)
    add_p(doc,'RAG PoC：从问题定位到多路检索验证',17,'637083',False,align=WD_ALIGN_PARAGRAPH.CENTER,after=30)
    add_p(doc,'[姓名占位]｜[学校 / 专业]｜AI 产品实习',12,'172B4D',False,align=WD_ALIGN_PARAGRAPH.CENTER,after=20)
    add_p(doc,'个人学习 PoC｜示例知识来自公开客服资料｜非宜家官方项目',10,'637083',False,align=WD_ALIGN_PARAGRAPH.CENTER,after=22)
    add_p(doc,'核心展示：发现问题 → 分层定位 → 设计实验 → 验证方案 → 评估效果、延迟与成本权衡',12,'102A56',True,align=WD_ALIGN_PARAGRAPH.CENTER,after=6)
    page_break(doc)
    # page helper content
    add_p(doc,'01｜用户、B2B 角色与业务问题',16,'2E74B5',True,style='Heading 1')
    add_p(doc,'定位：面向客服人员的 AI Copilot，而不是面向消费者的聊天机器人。',12,'102A56',True)
    tbl=doc.add_table(rows=4,cols=2); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl.style='Table Grid'
    rows=[('使用者','一线客服：查找政策、判断适用条件、组织回复。'),('管理者','客服主管 / 运营负责人：关注异常案例、规则一致性和处理效率。'),('决策者','客服中心负责人 / 数字化负责人：评估是否值得投入和扩大试点。'),('企业价值假设','减少查政策时间，降低规则判断错误，提高客服处理效率。')]
    for i,(a,b) in enumerate(rows):
        c1,c2=tbl.rows[i].cells; c1.width=Inches(1.55); c2.width=Inches(4.75); set_cell_shading(c1,'EAF2FF'); c1.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; c2.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_run(c1.paragraphs[0].add_run(a),11,'102A56',True); set_run(c2.paragraphs[0].add_run(b),11)
    page_break(doc)
    add_p(doc,'02｜产品方案与核心交互',16,'2E74B5',True,style='Heading 1'); add_p(doc,'AI 只提供有依据的建议回复；是否发送由人工客服决定。',12,'102A56',True); add_bullets(doc,['客户问题进入客服工作台','AI 检索企业知识库并生成建议回复','展示知识依据与处理状态','客服可采用、编辑、重新生成或转人工']); add_pic(doc,paths['workbench'],6.2); page_break(doc)
    add_p(doc,'03｜最初的 RAG 架构',16,'2E74B5',True,style='Heading 1'); add_p(doc,'初始链路先验证“企业政策文档 → 知识检索 → LLM 回答”能否跑通。',11); add_pic(doc,paths['initial'],6.25); add_p(doc,'核心组件：Chunk、Embedding、Hybrid Search、Rerank、Top K。初始风险是 Chunk 质量与知识库外误召回。',11,'172B4D'); page_break(doc)
    add_p(doc,'04｜实验 1：Chunk 粒度问题（V1）',16,'2E74B5',True,style='Heading 1'); add_p(doc,'原始 PDF 直接入库后，普通换行切分会形成没有完整语义的碎片；双换行切分又可能产生 700–1000 字的大块。',11); add_bullets(doc,['大 Chunk：上下文更完整，但混入多主题、导航、URL 和网页噪声。','过碎 Chunk：标题与答案分离，单独使用没有完整含义。','实验结论：目标不是固定长度，而是完整、单一的语义单元。']); page_break(doc)
    add_p(doc,'05｜实验 2：FAQ 级 Chunk（V2）',16,'2E74B5',True,style='Heading 1'); add_p(doc,'将公开客服资料清洗为 Markdown，尽量做到“一个问题 / 一条规则 ≈ 一个 Chunk”。',11); add_bullets(doc,['保留退货期限、商品排除范围、商品状态、会员卡与配送规则。','删除网页导航、URL、备案及推广信息。','单一意图问题更容易命中对应规则，检索结果更短、更干净。']); add_pic(doc,paths['experiments'],6.15); page_break(doc)
    add_p(doc,'06｜V1 与 V2：固定测试集与观察',16,'2E74B5',True,style='Heading 1'); add_p(doc,'测试方法：固定测试集 → 每轮使用相同问题 → 检查 Top1 / Top3 / 噪声 / 最终回答 → 记录异常案例。',11,'102A56',True)
    tbl=doc.add_table(rows=6,cols=3); tbl.style='Table Grid'; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    data=[('测试类型','V1 观察','V2 观察'),('单一 FAQ','能找到，但大 Chunk 噪声较多','FAQ 级规则更精准'),('手机号 / 配送改期','答案埋在大 Chunk 中','对应规则更容易 Top1'),('复合规则','大 Chunk 有时更完整','可能遗漏第二条必要规则'),('知识库外问题','仍会召回最相似内容','仍需 LLM 拒答'),('实验限制','配置需核验','不将变化全部归因于 Chunk')]
    for r,row in enumerate(data):
        for c,v in enumerate(row):
            cell=tbl.rows[r].cells[c]; cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r==0: set_cell_shading(cell,'F2F4F7')
            set_run(cell.paragraphs[0].add_run(v),10,'102A56' if r==0 else '172B4D',r==0)
    page_break(doc)
    add_p(doc,'07｜复合问题失败案例',16,'2E74B5',True,style='Heading 1'); add_p(doc,'V2 的关键短板不是单一规则，而是跨 Chunk 的证据完整性。',12,'102A56',True)
    add_p(doc,'案例 A：我是会员，半年前买了一台家电，可以退吗？',12,'172B4D',True); add_p(doc,'需要同时获得“会员 365 天规则”与“家用电器属于排除范围”。V2 的单次检索主要召回前者。',11)
    add_p(doc,'案例 B：商品在退货期限内，但已经使用过了，可以退吗？',12,'172B4D',True); add_p(doc,'需要同时获得“退货期限”与“货品完好 / 无使用痕迹”。V2 未稳定召回第二条规则。',11)
    add_p(doc,'证据说明：基于测试记录整理；缺失的原始 V2 截图未使用模拟界面替代。',10,'637083',False); page_break(doc)
    add_p(doc,'08｜问题定位方法与 Top K 实验',16,'2E74B5',True,style='Heading 1'); add_pic(doc,paths['diagnosis'],6.25); add_p(doc,'Top K 从 3 提到 5 后，两道复合题仍未稳定召回第二条规则。因此问题不只是“找得不够多”，而是单次检索难以兼顾多个独立意图。',11,'172B4D'); page_break(doc)
    add_p(doc,'09｜实验 3：复合 Query 的多路检索（V2.1）',16,'2E74B5',True,style='Heading 1'); add_p(doc,'将复杂问题拆为两个独立子问题，分别检索，再合并两路 Evidence 交给最终 LLM。',11); add_pic(doc,paths['v21'],6.25); page_break(doc)
    add_p(doc,'10｜两个复合问题成功案例',16,'2E74B5',True,style='Heading 1'); add_p(doc,'真实 Dify 运行截图：两题均经过查询拆解、双路知识检索、合并结果和最终 LLM 回答。',11,'102A56',True)
    for label,path in [('会员 + 半年 + 家电：正确识别家电排除范围',paths['success_appliance']),('期限内 + 已使用：正确识别“无使用痕迹”条件',paths['success_used'])]:
        add_p(doc,label,11,'172B4D',True); add_pic(doc,path,4.7)
    page_break(doc)
    add_p(doc,'11｜效果、延迟与成本 Trade-off',16,'2E74B5',True,style='Heading 1'); add_p(doc,'V2.1 的收益不是“全面更好”，而是针对复合问题用更多调用换取更完整的证据。',12,'102A56',True)
    add_bullets(doc,['收益：两个复合问题均取得正确结论。','延迟观察：会员 + 家电约 14 秒；期限内 + 已使用约 17 秒。','成本影响：额外的查询拆解 LLM 与第二次检索增加调用量。','产品判断：复杂问题可走多路检索；后续需继续评估是否按问题复杂度动态路由。']); page_break(doc)
    add_p(doc,'12｜高保真产品原型：客服工作台与知识依据',16,'2E74B5',True,style='Heading 1'); add_pic(doc,paths['sources'],6.25); page_break(doc)
    add_p(doc,'13｜高保真产品原型：拒答、转人工与主管视角',16,'2E74B5',True,style='Heading 1'); add_pic(doc,paths['escalation'],6.1); add_pic(doc,paths['admin'],6.1); page_break(doc)
    add_p(doc,'14｜企业价值假设与待验证指标',16,'2E74B5',True,style='Heading 1'); add_p(doc,'以下是上线后需要验证的指标，不代表当前 PoC 的真实业务结果。',11,'C2414B',True); add_bullets(doc,['AI 建议回复采纳率','人工修改率','平均处理时长','错误回答率','正确拒答率','客服满意度','单次 AI 调用成本']); page_break(doc)
    add_p(doc,'15｜项目限制与下一步',16,'2E74B5',True,style='Heading 1'); add_bullets(doc,['当前是个人 PoC，不含真实客服用户、生产流量或 ROI 验证。','知识来自公开客服资料，仅用于学习与产品能力展示。','下一步：补全真实测试截图，验证真实使用流程与人工反馈闭环。','下一步：评估复杂问题动态路由、延迟优化与成本控制。']); page_break(doc)
    add_p(doc,'附录｜90 秒 Demo 分镜、旁白与字幕稿',16,'2E74B5',True,style='Heading 1')
    script=[('0–10 秒','问题与定位','“这是一个面向一线客服的企业 AI 知识助手。它基于公开客服政策资料，为客服生成有依据的建议回复。”'),('10–25 秒','简单问题','“先看单一问题：会员卡丢失怎么办？系统直接命中对应 FAQ，并展示知识依据，客服可以采用或编辑。”'),('25–55 秒','复合问题','“再看会员、半年、家电这类复合问题。系统先拆成两个子问题，分别检索退货期限和家电例外，再合并证据。”'),('55–68 秒','拒答场景','“如果用户询问当前知识库没有的门店营业时间，系统不会编造答案，而是提示知识依据不足并支持转人工。”'),('68–82 秒','实验洞察','“实验发现：FAQ 级 Chunk 提升单一意图精准度，但复合问题可能漏规则。多路检索改善效果，但带来延迟和成本。”'),('82–90 秒','结尾','“这个 PoC 的重点不是替代客服，而是帮助客服更快、更有依据地完成判断。”')]
    tbl=doc.add_table(rows=1,cols=3); tbl.style='Table Grid'; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,v in enumerate(['时间','画面','旁白 / 字幕']): set_cell_shading(tbl.rows[0].cells[i],'F2F4F7'); set_run(tbl.rows[0].cells[i].paragraphs[0].add_run(v),10,'102A56',True)
    for a,b,c in script:
        cells=tbl.add_row().cells
        for cell,v in zip(cells,[a,b,c]): set_run(cell.paragraphs[0].add_run(v),9.5)
    out=ROOT/'企业AI客服RAG_PoC_作品集.docx'; doc.save(out); return out

def main():
    initial=architecture_initial(); v21=architecture_v21(); exp=experiments(); diag=diagnosis(); wb=prototype_workbench(); src=prototype_sources(); esc=prototype_escalation(); adm=prototype_admin()
    success_used=Path(r"C:\Users\Gloria\AppData\Local\Temp\codex-file-preview-gXZDTO\image.png")
    success_appliance=Path(r"C:\Users\Gloria\AppData\Local\Temp\codex-file-preview-5QlCpI\image.png")
    paths={'initial':initial,'v21':v21,'experiments':exp,'diagnosis':diag,'workbench':wb,'sources':src,'escalation':esc,'admin':adm,'success_used':success_used,'success_appliance':success_appliance}
    long_image(paths); docx(paths)
    (ROOT/'90秒Demo脚本.txt').write_text('企业 AI 客服知识助手｜90 秒 Demo\n\n0–10秒：介绍 B2B 客服 Copilot 与公开资料模拟场景。\n10–25秒：演示“会员卡丢失”单一 FAQ，展示知识依据。\n25–55秒：演示“会员+半年+家电”，展示查询拆解、双路检索与合并证据。\n55–68秒：演示知识库外问题拒答与转人工。\n68–82秒：说明三轮实验与效果/延迟/成本权衡。\n82–90秒：总结：AI 辅助人工，让客服更快且更有依据地判断。\n',encoding='utf-8')
    print('DONE', ROOT)

if __name__=='__main__': main()
